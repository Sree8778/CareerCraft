# backend/document_generator.py
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from weasyprint import HTML
from jinja2 import Environment, FileSystemLoader
import os
from bs4 import BeautifulSoup

# --- Helper function to clean up extra whitespace (more robust) ---
def clean_text(text: str) -> str:
    """
    Removes excessive whitespace, including multiple newlines, spaces, and tabs,
    and consolidates them to single spaces or newlines as appropriate.
    Also strips leading/trailing whitespace.
    """
    if not text:
        return ""
    # Replace multiple spaces/tabs with a single space
    cleaned_text = re.sub(r'[ \t]+', ' ', text)
    # Replace two or more newlines with a single newline
    cleaned_text = re.sub(r'\n{2,}', '\n', cleaned_text)
    # Trim each line
    cleaned_text = '\n'.join([line.strip() for line in cleaned_text.split('\n')])
    return cleaned_text.strip()


# --- Helper to add HTML content to a DOCX paragraph ---
def add_html_to_docx_paragraph(doc, paragraph, html_content, normal_style=None):
    if not html_content:
        return

    # Clean the HTML content first to remove excessive internal whitespace before parsing
    # This specifically target spaces/newlines *within* the HTML string
    html_content = re.sub(r'>\s+<', '><', html_content) # Remove whitespace between tags
    html_content = re.sub(r'\s{2,}', ' ', html_content) # Consolidate multiple spaces within text

    soup = BeautifulSoup(html_content, 'html.parser')

    # Iterate through direct children of the soup body or the entire soup if no body
    for element in soup.contents if soup.body else soup.contents:
        if element.name == 'ul':
            for li in element.find_all('li'):
                # Start a new paragraph for each list item
                li_p = paragraph.insert_paragraph_before() if paragraph.text else paragraph
                li_p.style = normal_style # Apply normal style to list item
                
                # Set bullet style (e.g., Unicode bullet)
                li_p.style.list_style = True # This doesn't actually apply a list style by itself, needs numbering.
                # A proper bullet list needs a numbering definition. For simplicity, just add a symbol and indent.
                li_p.add_run('•\t') # Add a simple bullet char and tab
                li_p.paragraph_format.left_indent = Pt(36) # Indent the text
                li_p.paragraph_format.first_line_indent = Pt(-18) # Hang indent for bullet

                # Process inline HTML within the list item
                process_inline_html_for_docx(li_p, str(li.decode_contents()), normal_style)
                paragraph = doc.add_paragraph() # Prepare for next content after list

        elif element.name == 'p':
            p_tag = paragraph if not paragraph.text else doc.add_paragraph(style=normal_style) # Use existing or new p
            process_inline_html_for_docx(p_tag, str(element.decode_contents()), normal_style)
            paragraph = doc.add_paragraph() # Add a new paragraph for separation

        elif element.name == 'br':
            paragraph.add_break() # Add line break
        
        elif element.string:
            run = paragraph.add_run(str(element.string).strip())
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        else: # Handle other inline elements or direct text content not inside p/ul
            process_inline_html_for_docx(paragraph, str(element), normal_style)


# Helper for recursive processing of inline HTML for DOCX
def process_inline_html_for_docx(parent_obj, html_snippet, normal_style):
    temp_soup = BeautifulSoup(html_snippet, 'html.parser')
    for content in temp_soup.contents:
        if content.name == 'strong':
            run = parent_obj.add_run(content.get_text())
            run.bold = True
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        elif content.name == 'em':
            run = parent_obj.add_run(content.get_text())
            run.italic = True
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        elif content.name == 'a': # Handle links
            run = parent_obj.add_run(content.get_text())
            run.font.underline = True
            run.font.color.rgb = RGBColor(0x0000FF) # Blue color for links
            # Link functionality itself is complex and usually requires a separate relationship in DOCX
            # For simplicity, just style the text.
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        elif content.name == 'br':
            parent_obj.add_break()
        elif content.string:
            run = parent_obj.add_run(str(content.string).strip()) # Trim string content
            if normal_style:
                run.font.name = normal_style.font.name
                run.font.size = normal_style.font.size
        elif content.contents: # Recursively process nested tags
            process_inline_html_for_docx(parent_obj, str(content), normal_style)


# --- DOCX GENERATION ---
def generate_docx_from_data(data):
    doc = Document()
    style = data.get('styleOptions', {})
    font_name = style.get('fontFamily', 'Calibri').split(',')[0]
    font_size = style.get('fontSize', 11)
    accent_color_hex = style.get('accentColor', '#34495e').lstrip('#')
    accent_color_rgb = RGBColor.from_string(accent_color_hex)

    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size)

    try:
        heading_style = doc.styles['SectionHeading']
    except KeyError:
        heading_style = doc.styles.add_style('SectionHeading', 1)
    heading_style.font.name = font_name
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.font.color.rgb = accent_color_rgb
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)

    # --- Header ---
    personal = data.get('personal', {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(personal.get('name', ''))
    runner.bold = True
    runner.font.name = font_name
    runner.font.size = Pt(24)
    runner.font.color.rgb = accent_color_rgb

    contact_items = [
        personal.get('email'),
        personal.get('phone'),
        personal.get('location')
    ]
    if personal.get('legalStatus') and personal.get('legalStatus') != 'Prefer not to say':
        contact_items.append(personal.get('legalStatus'))

    contact_info = " | ".join(filter(None, contact_items))
    p = doc.add_paragraph(contact_info)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Add a blank line for spacing

    # --- Sections (now parsing HTML from AI) ---
    if data.get('summary'):
        doc.add_paragraph('Summary', style='SectionHeading')
        p = doc.add_paragraph(style=normal_style)
        add_html_to_docx_paragraph(doc, p, data.get('summary'), normal_style)


    if data.get('experience') and any(e.get('jobTitle') for e in data.get('experience')):
        doc.add_paragraph('Experience', style='SectionHeading')
        for exp in data['experience']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(exp.get('jobTitle', '')).bold = True
            p.add_run(f"\n{exp.get('company', '')} | {exp.get('dates', '')}\n").italic = True
            add_html_to_docx_paragraph(doc, p, exp.get('description', ''), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get('education') and any(e.get('degree') for e in data.get('education')):
        doc.add_paragraph('Education', style='SectionHeading')
        for edu in data['education']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(edu.get('degree', '')).bold = True
            p.add_run(f", {edu.get('institution', '')}\n")
            gpa_part = (f' | GPA: {edu.get("gpa")}' if edu.get('gpa') else '')
            p.add_run(f"{edu.get('graduationYear', '')}{gpa_part}\n").italic = True
            add_html_to_docx_paragraph(doc, p, edu.get('achievements', ''), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get('skills') and any(s.get('skills_list') for s in data.get('skills')):
        doc.add_paragraph('Skills', style='SectionHeading')
        for skill in data['skills']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(f"{skill.get('category', '')}: ").bold = True
            # For skills_list, apply clean_text directly as it's plain text from AI
            p.add_run(clean_text(skill.get('skills_list', ''))) 
            p.paragraph_format.space_after = Pt(6)

    if data.get('projects') and any(p.get('title') for p in data.get('projects')):
        doc.add_paragraph('Projects', style='SectionHeading')
        for proj in data['projects']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(proj.get('title', '')).bold = True
            p.add_run(f" ({proj.get('date', '')})\n").italic = True
            add_html_to_docx_paragraph(doc, p, proj.get('description', ''), normal_style)
            p.paragraph_format.space_after = Pt(12)

    if data.get('publications') and any(pub.get('title') for pub in data.get('publications')):
        doc.add_paragraph('Publications', style='SectionHeading')
        for pub in data['publications']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(pub.get('title', '')).bold = True
            p.add_run(f" ({pub.get('date', '')})\n").italic = True
            p.add_run(f"{pub.get('authors', '')} - {pub.get('journal', '')}\n").font.size = Pt(font_size - 1)
            if pub.get('link'):
                p.add_run(f"Link: {pub.get('link')}").font.size = Pt(font_size - 1)
            p.paragraph_format.space_after = Pt(12)


    if data.get('certifications') and any(c.get('name') for c in data.get('certifications')):
        doc.add_paragraph('Certifications', style='SectionHeading')
        for cert in data['certifications']:
            p = doc.add_paragraph(style=normal_style)
            p.add_run(cert.get('name', '')).bold = True
            
            issuer_date_text = cert.get('issuer', '')
            if cert.get('date'):
                issuer_date_text += f" | {cert.get('date')}"
            
            p.add_run(f"\n{issuer_date_text}").italic = True
            p.paragraph_format.space_after = Pt(12)
            
    return doc


# --- PDF GENERATION ---
def generate_pdf_from_data(data):
    # Ensure data values are cleaned before rendering, especially for plain text fields
    # HTML fields should generally be passed as-is to Jinja2 and WeasyPrint to preserve structure.
    # We apply a light clean on relevant fields directly in the template render call now,
    # or rely on the HTML/CSS for formatting.
    env = Environment(loader=FileSystemLoader(os.path.join(os.path.dirname(__file__), 'assets')))
    template = env.get_template('resume_template.html')
    
    # Pre-clean skills_list if it's treated as plain text, before passing to template
    if data.get('skills'):
        for skill in data['skills']:
            skill['skills_list'] = clean_text(skill.get('skills_list', ''))

    rendered_html = template.render(**data)
    
    return HTML(string=rendered_html).write_pdf()